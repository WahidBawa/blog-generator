import os
from openai import OpenAI
from functools import lru_cache
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SYSTEM_PROMPT = """
[As a system, you are tasked with writing well-informed blog posts.]

You will receive a topic from a user for each blog post. Write each blog post to be well-researched, thorough in length, and informative. The content should be written in an approachable and easy-to-follow manner, while still maintaining depth and conveying expert knowledge.

Blog post titles will be provided in subsequent messages, and you will generate content accordingly.

# Output Format

Provide the blog post in a clear, structured format with appropriate headings to enhance readability and aid in content navigation. Ensure the language is both engaging and authoritative.

These following is what you should use to clearly format what the text should look like. For instance title would always have a Heading 1 centered format. Ensure that you ONLY use these following symbols in your formatting.

# for Heading 1 (centered).
## for Heading 2.
### for Heading 3.
* for bullet points.
Other text is treated as a paragraph.
"""

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")


@lru_cache()
def openapi_client():
    client = OpenAI(api_key=OPENAI_API_KEY)

    return client


def generate_blog(message: str) -> str | None:
    completion = openapi_client().chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": message,
            },
        ],
    )

    return completion.choices[0].message.content


def set_paragraph_font(paragraph, font_name="Times New Roman", font_size=12):
    """Set font for a given paragraph."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)


def add_text_with_formatting(paragraph, text):
    """Add text to a paragraph, handling bolded sections marked with **."""
    while "**" in text:
        before_bold, rest = text.split("**", 1)
        bold_text, after_bold = rest.split("**", 1)

        if before_bold:
            run = paragraph.add_run(before_bold)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        bold_run = paragraph.add_run(bold_text)
        bold_run.font.name = "Times New Roman"
        bold_run.font.size = Pt(12)
        bold_run.bold = True

        text = after_bold

    if text:
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def write_to_docx(input_text, output_file):
    doc = Document()

    lines = input_text.splitlines()

    for line in lines:
        line = line.strip()

        if line.startswith("# "):
            paragraph = doc.add_heading(level=1)
            paragraph.text = line[2:]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_font(paragraph, font_name="Times New Roman", font_size=16)

        elif line.startswith("## "):
            paragraph = doc.add_heading(level=2)
            paragraph.text = line[3:]
            set_paragraph_font(paragraph, font_name="Times New Roman", font_size=14)

        elif line.startswith("### "):
            paragraph = doc.add_heading(level=3)
            paragraph.text = line[4:]
            set_paragraph_font(paragraph, font_name="Times New Roman", font_size=12)

        elif line.startswith("* "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_text_with_formatting(paragraph, line[2:])

        elif line.startswith("• "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_text_with_formatting(paragraph, line[2:])

        elif line:
            paragraph = doc.add_paragraph()
            add_text_with_formatting(paragraph, line)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(12)

        else:
            doc.add_paragraph("")

    styles = doc.styles
    normal_style = styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Times New Roman"
    normal_font.size = Pt(12)

    doc.save(output_file)
    print(f"Document saved as {output_file}")
